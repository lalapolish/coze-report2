import re
import io
import os
import uuid
import pandas as pd
import matplotlib.pyplot as plt
from fastapi import FastAPI, Request
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from typing import Optional
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = FastAPI(openapi_version="3.0.0")

if not os.path.exists("static"):
    os.makedirs("static")
app.mount("/static", StaticFiles(directory="static"), name="static")

class ReportInput(BaseModel):
    ch2_text: Optional[str] = ""
    ch3_text: Optional[str] = ""
    ch4_text: Optional[str] = ""
    ch5_text: Optional[str] = ""
    ch6_text: Optional[str] = ""
    ch7_text: Optional[str] = ""

# --- 绘图配置 ---
try:
    plt.rcParams['font.sans-serif'] = ['SimHei', 'DejaVu Sans', 'Arial Unicode MS']
except:
    pass
plt.rcParams['axes.unicode_minus'] = False
BLUE_COLOR = 'royalblue'

def clean_data(text):
    if pd.isna(text): return 0
    val = str(text).replace('$', '').replace(',', '').replace('%', '').replace('*', '').strip()
    try:
        if '.' in val: return float(val)
        return int(val)
    except:
        return val

def md_table_to_df(md_text):
    # 允许更灵活的表格解析
    lines = [l.strip() for l in md_text.strip().split('\n') if '|' in l]
    if len(lines) < 2: return None
    headers = [re.sub(r'[\$\*]', '', c).strip() for c in lines[0].split('|') if c.strip()]
    data = []
    for line in lines:
        if '---' in line or line == lines[0]: continue
        row = [clean_data(c) for c in line.split('|') if c.strip()]
        if len(row) >= len(headers):
            data.append(row[:len(headers)])
    return pd.DataFrame(data, columns=headers) if data else None

def generate_chart(df, title, fig_no):
    plt.figure(figsize=(9, 5))
    img_stream = io.BytesIO()
    try:
        x_data = df.iloc[:, 0].astype(str)
        y_data = pd.to_numeric(df.iloc[:, 1], errors='coerce').fillna(0)
        
        # 自动判定：第4、5章或标题带占比/分布画饼图，其他画柱状图
        if fig_no in [4, 5] or any(kw in title for kw in ["分布", "占比", "结构"]):
            plt.pie(y_data, labels=x_data, autopct='%1.1f%%', colors=plt.cm.Pastel1.colors)
        else:
            bars = plt.bar(x_data, y_data, color=BLUE_COLOR, width=0.6)
            for bar in bars:
                plt.text(bar.get_x() + bar.get_width()/2, bar.get_height(), f'{bar.get_height()}', ha='center', va='bottom')
        
        plt.title(title, fontsize=12, pad=15)
        plt.tight_layout()
        plt.savefig(img_stream, format='png', dpi=200)
    finally: 
        plt.close()
    img_stream.seek(0)
    return img_stream

def set_style(obj, is_title=False):
    if hasattr(obj, 'runs'):
        for run in obj.runs:
            run.font.size = Pt(14 if is_title else 12)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def process_content(doc, full_text):
    # --- 【关键改进】预处理：修复大模型格式粘连 ---
    # 1. 在表格起始符 | 前强制换行（如果前面不是换行的话）
    full_text = re.sub(r'([^\n])\s*\|', r'\1\n|', full_text)
    # 2. 在“图 X:”或“表 X:”前强制换行
    full_text = re.sub(r'([^\n])\s*(\*\*?[图表]\s?\d+[:：])', r'\1\n\2', full_text)
    
    # 使用更强大的切割逻辑，将标题、表格、普通文本分开
    parts = re.split(r'(##+ .*?\n|(?:\*\*?)?[图表]\s?\d+[:：].*?\n|(?:\n|^)\|[\s\S]*?\|(?:\n|$))', full_text)
    
    current_fig_title = None
    
    for part in parts:
        if not part or not part.strip(): continue
        part_s = part.strip()
        
        # 1. 处理二级或三级标题
        if part_s.startswith('##'):
            h = doc.add_heading(part_s.replace('#','').strip(), level=2)
            set_style(h, True)
            
        # 2. 处理图表标题行
        elif re.match(r'(\*\*?)?[图表]\s?\d+[:：]', part_s):
            current_fig_title = part_s.replace('*', '').strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(current_fig_title)
            run.bold = True
            set_style(p)
            
        # 3. 处理表格和图表转换
        elif part_s.startswith('|'):
            df = md_table_to_df(part_s)
            if df is not None:
                # 只有标题里带“图”字才转成图片
                if current_fig_title and "图" in current_fig_title:
                    try:
                        fig_no_match = re.search(r'\d+', current_fig_title)
                        fig_no = int(fig_no_match.group()) if fig_no_match else 0
                        img = generate_chart(df, current_fig_title, fig_no)
                        doc.add_picture(img, width=Inches(5.5))
                        # 图片居中
                        last_p = doc.paragraphs[-1]
                        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        print(f"绘图失败: {e}")
                    current_fig_title = None 
                else:
                    # 否则插入普通 Word 表格
                    table = doc.add_table(rows=1, cols=len(df.columns))
                    table.style = 'Table Grid'
                    for i, col in enumerate(df.columns): table.rows[0].cells[i].text = str(col)
                    for _, row in df.iterrows():
                        row_cells = table.add_row().cells
                        for i, val in enumerate(row): row_cells[i].text = str(val)
            
        # 4. 处理普通文本段落
        else:
            # 过滤掉孤立的 Markdown 标记
            clean_text = part_s.replace('$', '').replace('***', '').replace('---', '')
            if len(clean_text.strip()) > 0:
                p = doc.add_paragraph(clean_text)
                set_style(p)

@app.post("/generate_report_word")
async def generate_report_word(input_data: ReportInput, request: Request):
    try:
        doc = Document()
        for i in range(2, 8):
            content = getattr(input_data, f"ch{i}_text", "")
            if content and len(content.strip()) > 5:
                process_content(doc, content)
        
        file_id = uuid.uuid4().hex[:8]
        file_name = f"report_{file_id}.docx"
        file_path = os.path.join("static", file_name)
        doc.save(file_path)
        
        base_url = str(request.base_url).rstrip('/')
        return {
            "status": "success",
            "file_url": f"{base_url}/static/{file_name}",
            "message": "文档生成完毕，请下载"
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=int(os.environ.get("PORT", 8080)))
